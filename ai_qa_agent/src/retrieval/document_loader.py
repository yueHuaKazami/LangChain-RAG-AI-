"""通用文档加载器。

支持格式：.md / .txt / .pdf / .docx / .doc
- md、txt：Python 内置，直接读取文本
- pdf：pypdf 解析
- docx：python-docx 解析
- doc：pywin32 调用 Word COM 解析（仅 Windows + 已安装 Word）

每种格式采用延迟导入，未使用到的库不会影响模块加载。
统一输出 LangChain 标准 Document 对象，供向量库构建使用。
"""
import os
from typing import Callable

from langchain_core.documents import Document

from config import settings


def load_document(file_path: str) -> Document:
    """根据文件扩展名加载单个文档。

    Args:
        file_path: 文件绝对路径。

    Returns:
        包含文件正文和元数据的 Document 对象。

    Raises:
        ValueError: 文件格式不支持。
        FileNotFoundError: 文件不存在。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在：{file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    loader_map: dict[str, Callable[[str], str]] = {
        ".md": _load_text,
        ".txt": _load_text,
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".doc": _load_doc,
    }

    loader = loader_map.get(ext)
    if loader is None:
        raise ValueError(
            f"不支持的文件格式：{ext}。"
            f"支持的格式：{sorted(settings.SUPPORTED_EXTENSIONS)}"
        )

    text = loader(file_path)
    filename = os.path.basename(file_path)
    return Document(page_content=text, metadata={"source": filename})


def load_documents_from_dir(dir_path: str = settings.DATA_DIR) -> list[Document]:
    """遍历目录加载所有受支持格式的文档。

    按 文件名排序保证加载顺序稳定，便于调试和复现。

    Args:
        dir_path: 数据目录路径，默认为 settings.DATA_DIR。

    Returns:
        Document 对象列表。

    Raises:
        FileNotFoundError: 目录不存在或没有可加载的文档。
    """
    if not os.path.isdir(dir_path):
        raise FileNotFoundError(f"数据目录不存在：{dir_path}")

    documents: list[Document] = []
    for filename in sorted(os.listdir(dir_path)):
        ext = os.path.splitext(filename)[1].lower()
        if ext not in settings.SUPPORTED_EXTENSIONS:
            continue
        file_path = os.path.join(dir_path, filename)
        doc = load_document(file_path)
        documents.append(doc)
        print(f"[loaders] 已加载文档：{filename}（{len(doc.page_content)} 字符）")

    if not documents:
        raise FileNotFoundError(
            f"目录 {dir_path} 中没有找到可加载的文档。"
            f"支持格式：{sorted(settings.SUPPORTED_EXTENSIONS)}"
        )

    print(f"[loaders] 共加载 {len(documents)} 个文档")
    return documents


# === 各格式具体加载实现（私有函数，延迟导入第三方库）===


def _load_text(file_path: str) -> str:
    """加载纯文本文件（md / txt）。

    Args:
        file_path: 文件路径。

    Returns:
        文件文本内容。
    """
    with open(file_path, encoding="utf-8") as f:
        return f.read()


def _load_pdf(file_path: str) -> str:
    """加载 PDF 文件，提取全部文本。

    使用 pypdf 逐页提取并拼接，页与页之间以换行分隔。

    Args:
        file_path: PDF 文件路径。

    Returns:
        PDF 全部文本内容。
    """
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _load_docx(file_path: str) -> str:
    """加载 DOCX 文件，提取全部段落文本。

    使用 python-docx 读取段落，段落间以换行分隔。

    Args:
        file_path: DOCX 文件路径。

    Returns:
        DOCX 全部文本内容。
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _load_doc(file_path: str) -> str:
    """加载 DOC（老格式）文件。

    python-docx 不支持 .doc 二进制格式，此处通过 pywin32 调用
    Windows Word COM 组件解析。仅适用于 Windows + 已安装 Microsoft Word。

    Args:
        file_path: DOC 文件路径。

    Returns:
        DOC 全部文本内容。

    Raises:
        RuntimeError: 非 Windows 环境或未安装 Word，无法解析 .doc。
    """
    import sys

    if sys.platform != "win32":
        raise RuntimeError(
            f".doc 格式仅在 Windows 下通过 Word COM 解析，当前平台：{sys.platform}。"
            "请将 .doc 转换为 .docx 后重试。"
        )

    try:
        import win32com.client
    except ImportError as e:
        raise RuntimeError(
            "未安装 pywin32，无法解析 .doc 文件。"
            "请执行 `uv add pywin32`，或将 .doc 转换为 .docx 后重试。"
        ) from e

    # 通过 Word COM 打开文档，提取文本后务必关闭，避免 Word 进程残留
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        # 路径需为绝对路径，Word COM 不接受相对路径
        doc = word.Documents.Open(os.path.abspath(file_path))
        text = doc.Content.Text
        return text.strip()
    except Exception as e:
        raise RuntimeError(
            f"通过 Word COM 解析 .doc 失败：{e}。"
            "请确认已安装 Microsoft Word，或将 .doc 转换为 .docx 后重试。"
        ) from e
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
